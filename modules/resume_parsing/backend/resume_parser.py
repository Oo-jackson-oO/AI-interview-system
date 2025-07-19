import json
import os
from docx import Document
import PyPDF2
import io
from docx2txt import process as docx2txt_process
from openai import OpenAI

class ResumeParser:
    def __init__(self):
        self.client = OpenAI(
            api_key='QcGCOyVichfHetzkUDeM:AUoiqAJtarlstnrJMcTI',
            base_url='https://spark-api-open.xf-yun.com/v1/'
        )
        
    def save_resume_to_mock_interview(self, text):
        """将简历文本保存到 modules/Mock_interview/resume.txt"""
        try:
            # 获取当前文件的目录路径
            current_dir = os.path.dirname(os.path.abspath(__file__))
            # 构建目标路径：从 modules/resume_parsing/backend 到 modules/Mock_interview
            mock_interview_dir = os.path.join(current_dir, '..', '..', 'Mock_interview')
            # 确保目录存在
            os.makedirs(mock_interview_dir, exist_ok=True)
            # 目标文件路径
            resume_file_path = os.path.join(mock_interview_dir, 'resume.txt')
            
            # 写入文件（如果存在则覆盖）
            with open(resume_file_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            print(f"简历文本已保存到: {resume_file_path}")
            return True
        except Exception as e:
            print(f"保存简历文本到 Mock_interview 目录失败: {str(e)}")
            return False
        
    def extract_text_from_file(self, file):
        """从文件中提取文本（支持PDF、DOCX、DOC）"""
        # 获取文件名并处理可能的None情况
        filename = getattr(file, 'filename', None) or getattr(file, 'name', 'unknown')
        
        # 调试信息
        print(f"处理文件: {filename}")
        print(f"文件对象属性: {dir(file)}")
        
        # 更安全的文件扩展名检测
        if '.' in filename:
            file_extension = filename.lower().split('.')[-1]
            print(f"从文件名检测到扩展名: {file_extension}")
        else:
            # 如果没有扩展名，尝试通过文件内容判断
            file.seek(0)
            header = file.read(8)
            file.seek(0)
            
            print(f"文件头部字节: {header}")
            
            if header.startswith(b'%PDF'):
                file_extension = 'pdf'
                print("通过文件头部识别为PDF")
            elif header.startswith(b'PK'):
                file_extension = 'docx'
                print("通过文件头部识别为DOCX")
            elif header.startswith(b'\xd0\xcf\x11\xe0'):
                file_extension = 'doc'
                print("通过文件头部识别为DOC")
            else:
                raise ValueError(f"无法识别文件格式，文件名: {filename}，头部字节: {header}")
        
        # 提取文本
        if file_extension == 'pdf':
            text = self.extract_text_from_pdf(file)
        elif file_extension == 'docx':
            text = self.extract_text_from_docx(file)
        elif file_extension == 'doc':
            text = self.extract_text_from_doc(file)
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}，文件名: {filename}")
        
        # 保存简历文本到 Mock_interview 目录
        self.save_resume_to_mock_interview(text)
        
        return text
    
    def extract_text_from_pdf(self, file):
        """从PDF文件中提取文本"""
        try:
            # 确保文件指针在开始位置
            file.seek(0)
            # 读取PDF文件
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
            text = ""
            
            # 提取每一页的文本
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text += page_text + "\n"
            
            result = text.strip()
            return result if result else "PDF文件似乎为空或无法读取文本内容"
            
        except Exception as e:
            raise Exception(f"PDF文件读取失败: {str(e)}")
        
    def extract_text_from_docx(self, file):
        """从DOCX文档中提取文本"""
        try:
            # 确保文件指针在开始位置
            file.seek(0)
            doc = Document(file)
            full_text = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():  # 只添加非空段落
                    full_text.append(para.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            result = '\n'.join(full_text)
            
            # 如果没有提取到文本，尝试其他方法
            if not result.strip():
                # 重置文件指针
                file.seek(0)
                # 尝试用docx2txt作为备用方案
                try:
                    # 将文件内容读入BytesIO对象
                    file_content = file.read()
                    file_like = io.BytesIO(file_content)
                    result = docx2txt_process(file_like)
                except Exception as docx2txt_error:
                    result = f"无法提取DOCX文件内容，主要方法和备用方法都失败了：{str(docx2txt_error)}"
            
            return result if result.strip() else "DOCX文件似乎为空或无法读取内容"
            
        except Exception as e:
            # 如果Document方法失败，尝试docx2txt
            try:
                file.seek(0)
                file_content = file.read()
                file_like = io.BytesIO(file_content)
                result = docx2txt_process(file_like)
                return result if result.strip() else "无法从DOCX文件提取文本内容"
            except Exception as e2:
                raise Exception(f"DOCX文件读取失败: {str(e)}，备用方法也失败: {str(e2)}")
    
    def extract_text_from_doc(self, file):
        """从DOC文档中提取文本"""
        try:
            # 确保文件指针在开始位置
            file.seek(0)
            # 对于老版本的DOC文件，使用docx2txt处理
            text = docx2txt_process(file)
            result = text.strip() if text else ""
            return result if result else "DOC文件似乎为空，请尝试转换为DOCX或PDF格式"
        except Exception as e:
            raise Exception(f"DOC文件读取失败: {str(e)}")
        
    # 保持向后兼容性
    def extract_text_from_word(self, file):
        """从Word文档中提取文本（保持向后兼容）"""
        return self.extract_text_from_docx(file)
    
    def analyze_resume(self, text):
        """分析简历内容 - 流式输出"""
        evaluation_criteria = """
        请根据以下标准分析这份简历（忽视排版问题）：
        1. 技能匹配（40分）- 简历中是否列出清晰、具体且实用的技能？技能是否具备一定的广度与深度？是否包括技术类或管理类工具（如 `Excel`、`Python`、`ERP系统`、`办公自动化` 等）？
        2. 经验匹配（30分）- - 是否清晰列出相关经历？经历是否有条理、无时间断档？
        3. 行业匹配度（20分）- 简历中是否清楚说明有哪些工作经历？岗位职责是否具体？是否展示出对某一类能力的积累？
        4. 成长与兴趣（10分）- 简历中是否体现出清晰的职业目标与兴趣？
        5. 是否可疑的简历信息。（扣分）·信息是否逻辑不清、夸大、缺失时间等？

        请使用Markdown格式给出详细的分析报告和各项评分，包含以下结构：
        - 请先给出一句总评价，然后再分析各个板块得分及其原因。
        - 最后，给出总得分，提供改进的意见与建议。
        - 使用# ## ###等标题分级
        - 使用**粗体**和*斜体*强调重要内容
        - 使用- 或1. 2. 等列表格式
        - 使用> 引用重要提示
        - 使用`代码`标记技术术语
        - 适当使用emoji表情符号

        得分标准：80-90优，60-79良，40-59中。
        """

        try:
            response = self.client.chat.completions.create(
                model='generalv3.5',
                messages=[
                    {"role": "user", "content": f"简历内容：{text}\n\n{evaluation_criteria}"}
                ],
                stream=True  # 启用流式输出
            )
            return self._parse_streaming_response(response)
        except Exception as e:
            raise Exception(f"API调用失败: {str(e)}")
    
    def analyze_resume_stream(self, text):
        """分析简历内容 - 真正的流式输出生成器"""
        evaluation_criteria = """
        请根据以下标准分析这份简历（忽视排版问题）：
        1. 技能匹配（40分）- 简历中是否列出清晰、具体且实用的技能？技能是否具备一定的广度与深度？是否包括技术类或管理类工具（如 `Excel`、`Python`、`ERP系统`、`办公自动化` 等）？
        2. 经验匹配（30分）- - 是否清晰列出相关经历？经历是否有条理、无时间断档？
        3. 行业匹配度（20分）- 简历中是否清楚说明有哪些工作经历？岗位职责是否具体？是否展示出对某一类能力的积累？
        4. 成长与兴趣（10分）- 简历中是否体现出清晰的职业目标与兴趣？
        5. 是否可疑的简历信息。（扣分）·信息是否逻辑不清、夸大、缺失时间等？

        请使用Markdown格式给出详细的分析报告和各项评分，包含以下结构：
        - 请先给出一句总评价，然后再分析各个板块得分及其原因。
        - 最后，给出总得分，提供改进的意见与建议。
        - 使用# ## ###等标题分级
        - 使用**粗体**和*斜体*强调重要内容
        - 使用- 或1. 2. 等列表格式
        - 使用> 引用重要提示
        - 使用`代码`标记技术术语
        - 适当使用emoji表情符号

        得分标准：80-90优，60-79良，40-59中。
        """

        try:
            response = self.client.chat.completions.create(
                model='generalv3.5',
                messages=[
                    {"role": "user", "content": f"简历内容：{text}\n\n{evaluation_criteria}"}
                ],
                stream=True  # 启用流式输出
            )
            
            for chunk in response:
                if chunk.choices[0].delta.content is not None:
                    content = chunk.choices[0].delta.content
                    # 确保内容被正确编码
                    if content:
                        yield content.encode('utf-8').decode('utf-8')
                    
        except Exception as e:
            error_msg = f"分析失败: {str(e)}"
            yield error_msg.encode('utf-8').decode('utf-8')
    
    def _parse_streaming_response(self, response):
        """解析流式响应"""
        full_response = ""
        try:
            for chunk in response:
                if chunk.choices[0].delta.content is not None:
                    content = chunk.choices[0].delta.content
                    full_response += content
        except Exception as e:
            raise Exception(f"流式响应解析失败: {str(e)}")
        return full_response
    
    def chat_with_ai(self, messages, new_message):
        """与AI进行对话 - 流式输出"""
        try:
            response = self.client.chat.completions.create(
                model='generalv3.5',
                messages=messages + [{"role": "user", "content": new_message}],
                stream=True  # 启用流式输出
            )
            return self._parse_streaming_response(response)
        except Exception as e:
            raise Exception(f"API调用失败: {str(e)}")
    
    def chat_with_ai_stream(self, messages, new_message):
        """与AI进行对话 - 真正的流式输出生成器"""
        try:
            response = self.client.chat.completions.create(
                model='generalv3.5',
                messages=messages + [{"role": "user", "content": new_message}],
                stream=True  # 启用流式输出
            )
            
            for chunk in response:
                if chunk.choices[0].delta.content is not None:
                    content = chunk.choices[0].delta.content
                    # 确保内容被正确编码
                    if content:
                        yield content.encode('utf-8').decode('utf-8')
                    
        except Exception as e:
            error_msg = f"对话失败: {str(e)}"
            yield error_msg.encode('utf-8').decode('utf-8') 